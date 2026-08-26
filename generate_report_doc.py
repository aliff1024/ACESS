import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # Page setup - Normal Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Palette
    COLOR_PRIMARY = RGBColor(30, 27, 75)      # Deep Indigo / Navy #1E1B4B
    COLOR_SECONDARY = RGBColor(79, 70, 229)   # Indigo #4F46E5
    COLOR_ACCENT = RGBColor(13, 148, 136)     # Teal #0D9488
    COLOR_TEXT = RGBColor(31, 41, 55)         # Dark Gray #1F2937
    COLOR_MUTED = RGBColor(107, 114, 128)     # Muted Gray #6B7280

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = COLOR_MUTED

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_ACCENT

    def add_callout(text, title="NOTE / IMPROVEMENT HIGHLIGHT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F5F3FF")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="4F46E5"/>'
            f'<w:top w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(2)
        r_t = cp.add_run(f"📌 {title}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = COLOR_SECONDARY

        r_b = cp.add_run(text)
        r_b.font.size = Pt(10)
        r_b.font.color.rgb = COLOR_TEXT

        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(4)
        p_after.paragraph_format.space_after = Pt(4)

    def add_improvement_box(user_role):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F9FAFB")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="18" w:space="0" w:color="0D9488"/>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="E5E7EB"/>'
            f'<w:right w:val="single" w:sz="6" w:space="0" w:color="E5E7EB"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="E5E7EB"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(4)
        r_t = cp.add_run(f"✍️ Self-Improvement & Development Notes: {user_role.upper()}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = COLOR_ACCENT

        r_hint = cp.add_run("[Document your specific additions, bug fixes, UI enhancements, and feature refinements for this user tier here]\n\n"
                            "• Improvement 1: _____________________________________________________\n"
                            "• Improvement 2: _____________________________________________________\n"
                            "• Improvement 3: _____________________________________________________\n"
                            "• Tested Status: [  ] Verified on Localhost   [  ] Verified on Vercel Production")
        r_hint.font.size = Pt(10)
        r_hint.font.color.rgb = COLOR_MUTED

        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(4)
        p_after.paragraph_format.space_after = Pt(4)

    # ──────────────────────────────────────────────────────────────────────────
    # COVER / HEADER
    # ──────────────────────────────────────────────────────────────────────────
    add_title("ACESS Application Testing & System Specification Report")
    add_subtitle("Final Year Project (PSM) Development, Feature Matrix, WCAG 2.1/2.2 Accessibility Engine, and Quality Assurance Plan")

    add_callout(
        "This document serves as an evaluation and testing benchmark for the ACESS Accessible Learning Platform. "
        "It details the 3 user personas (Admin, Educator, Learner), comprehensive WCAG accessibility mappings, "
        "the educator automated compliance checking engine, automated load/E2E testing methodologies (k6 & Playwright), "
        "and a complete database state summary.",
        "EXECUTIVE SUMMARY & PURPOSE"
    )

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 1: USER ROLES & PLATFORM CAPABILITIES
    # ──────────────────────────────────────────────────────────────────────────
    add_h1("1. User Role Capabilities & Functional Architecture")
    doc.add_paragraph("The ACESS platform implements strict multi-tiered Role-Based Access Control (RBAC) across three distinct user roles. Each role is equipped with tailored dashboards, workflows, and administrative privileges.")

    # --- 1.1 Admin ---
    add_h2("1.1 Administrator Portal (`/admin`)")
    doc.add_paragraph("System administrators possess global governance over platform health, account security, and curriculum quality. The administrator portal includes:")
    
    admin_bullets = [
        ("Executive Command Center: ", "Visual telemetry for platform engagement, total enrollments, active courses, certificate claims, and system health."),
        ("User Management (`/admin/users`): ", "Comprehensive directory to view all learner and educator accounts, edit user profile metadata, change account roles, and activate or suspend accounts."),
        ("Detailed User Telemetry (`/admin/users/[id]`): ", "Inspect individual learners' module progress, quiz attempt logs, earned certificates, and declared accessibility presets without modifying account sessions."),
        ("Course Governance (`/admin/courses`): ", "Review, inspect, publish, archive, or approve educator-authored courses and manage system-wide foundational courses."),
        ("Instructor Application Approval (`/admin/instructor-applications`): ", "Review educator vetting requests with qualifications and sample portfolios; 1-click promotion to educator role."),
        ("Centralized Certificate Oversight (`/admin/certificates`): ", "Inspect all issued digital credentials, verify cryptographic reference codes, and manage revoked certificates."),
        ("Feedback & Inquiry Center (`/admin/contact-messages`): ", "Triage incoming contact messages, support tickets, and accessibility improvement suggestions."),
        ("Automated Report Generator (`/admin/reports`): ", "Export high-level analytical summaries, accessibility adoption statistics, and course completion metrics into PDF and printable tables."),
        ("Universal Search: ", "Top-bar global search capable of instantaneous filtering across users, courses, certificates, and system features.")
    ]
    for b_title, b_desc in admin_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = COLOR_TEXT
        r2 = bp.add_run(b_desc)
        r2.font.color.rgb = COLOR_TEXT

    add_improvement_box("Administrator")

    # --- 1.2 Educator ---
    add_h2("1.2 Educator Portal (`/educator`)")
    doc.add_paragraph("Educators are empowered with specialized course creation tools, real-time learner monitoring, and automated accessibility auditing. Key features include:")
    
    educator_bullets = [
        ("Course Studio & Builder (`/educator/courses`): ", "Modular curriculum builder supporting structured chapters, lessons, checkpoints, and multi-format lesson content."),
        ("Rich Content Editor with Accessibility Tools: ", "WYSIWYG TipTap editor supporting accessible typography, custom callouts, dyslexia-friendly tables, and integrated media uploads."),
        ("Automated Accessibility Compliance Engine: ", "Real-time auditing of lesson markup calculating a 0-100 score, categorizing accessibility into Good, Warning, or Critical bands."),
        ("Interactive H5P & Video Checkpoints: ", "Embed interactive self-hosted H5P activities and in-video knowledge check questions that pause video playback for student comprehension."),
        ("Student Progress Intelligence (`/educator/students`): ", "Cohort tracking displaying individual completion rates, average scores, and automatic flags for 'At-Risk' or 'Inactive' learners (>14 days inactive)."),
        ("Deep Course Analytics (`/educator/analytics`): ", "Engagement heatmaps, completion funnels, lesson drop-off detection, and quiz score distribution graphs."),
        ("Custom Certificate Designer (`/educator/certificates`): ", "Create course-specific certificate templates with custom signatures, institutional badges, and issue verified credentials to graduating students.")
    ]
    for b_title, b_desc in educator_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = COLOR_TEXT
        r2 = bp.add_run(b_desc)
        r2.font.color.rgb = COLOR_TEXT

    add_improvement_box("Educator")

    # --- 1.3 Learner ---
    add_h2("1.3 Learner Portal (`/learner`)")
    doc.add_paragraph("Learners experience a highly adaptive, distraction-free environment tailored to their individual neurodivergent or sensory profile:")
    
    learner_bullets = [
        ("Adaptive Learning Workspace (`/learner/courses/[id]`): ", "Flexible lesson presentation supporting Slideshow View, Chunked Micro-learning, and Continuous Scroll modes."),
        ("Personalized Accessibility Presets: ", "1-click activation of tailored configurations for Dyslexia, ADHD, Autism, Vision Impairment, or customized preferences."),
        ("Assistive Reading Toolkit: ", "Integrated browser Text-to-Speech (TTS) with speed adjustment, Reading Spotlight focus ruler, dyslexia-friendly fonts, and background tinting."),
        ("Distraction-Free Mode: ", "Full-screen clean layout removing sidebars, headers, and extraneous visual elements to maximize focus."),
        ("Interactive Assessments & Quizzes: ", "Accessible quizzes featuring clear answer choices, instant feedback, and secure server-side grading without answer key exposure."),
        ("Gamification & Motivation (`/learner/achievements`): ", "Learning streaks, experience points (XP), badge milestone unlocks, and visual progress level rings."),
        ("Verifiable Digital Credentials (`/learner/certificates`): ", "Downloadable PDF certificates with high-contrast layout, QR code validation, and a public verification link (`/verify/[code]`).")
    ]
    for b_title, b_desc in learner_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = COLOR_TEXT
        r2 = bp.add_run(b_desc)
        r2.font.color.rgb = COLOR_TEXT

    add_improvement_box("Learner")

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 2: ACCESSIBILITY FEATURES & WCAG MAPPING
    # ──────────────────────────────────────────────────────────────────────────
    add_h1("2. In-Depth Accessibility Architecture & WCAG Compliance")
    doc.add_paragraph(
        "ACESS is designed from the ground up around the Web Content Accessibility Guidelines (WCAG 2.1 & 2.2 AA/AAA). "
        "The platform provides cognitive and sensory adaptations mapped directly to specific neurodivergent learning profiles."
    )

    # Preset Matrix Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Preset Name", "Target Learner Profile", "Active Adaptations Applied", "WCAG 2.1 / 2.2 Criteria"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1E1B4B")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    preset_rows = [
        (
            "Dyslexia Preset",
            "Learners with dyslexia, visual tracking difficulties, and reading fatigue.",
            "• Atkinson Hyperlegible / OpenDyslexic font (19px)\n"
            "• Line spacing 1.7×, Word spacing +40% (0.16em)\n"
            "• Cream background tint (#FFFDD0)\n"
            "• Text-to-Speech (TTS) enabled\n"
            "• Reading Spotlight focus ruler\n"
            "• Chunked content mode",
            "• 1.4.12 Text Spacing (AA)\n"
            "• 1.4.8 Visual Presentation (AAA)\n"
            "• 1.2.1 Audio Description (A)\n"
            "• 2.4.7 Focus Visible (AA)"
        ),
        (
            "ADHD Preset",
            "Learners with attention deficit, executive dysfunction, and working memory load.",
            "• Chunked Micro-learning (1 concept at a time)\n"
            "• Distraction-Free full-screen mode\n"
            "• Reading Spotlight ruler\n"
            "• Interactive Step-by-Step Task Checklist\n"
            "• Visual Progress Timeline & Checkpoints\n"
            "• Soft Grey background tint (#F3F4F6)\n"
            "• Auto-save visual indicator",
            "• 2.2.2 Pause, Stop, Hide (A)\n"
            "• 3.1.5 Reading Level (AAA)\n"
            "• 3.3.2 Labels or Instructions (A)\n"
            "• 2.2.1 Timing Adjustable (A)"
        ),
        (
            "Autism & Sensory",
            "Learners on the autism spectrum with sensory sensitivity or unpredictability anxiety.",
            "• Predictable Continuous Scroll layout\n"
            "• Muted pastel color palette (#E0F2FE pale blue)\n"
            "• Zero animations (data-animation-level='none')\n"
            "• Visual Schedule & step-by-step guidance\n"
            "• Reduced sensory visual crowding\n"
            "• Distraction-Free mode enabled",
            "• 2.3.3 Animation from Interactions (AAA)\n"
            "• 3.2.3 Consistent Navigation (AA)\n"
            "• 1.4.3 Contrast (Minimum) (AA)\n"
            "• 3.3.5 Help (AAA)"
        ),
        (
            "Vision Impairment",
            "Learners with low vision, contrast sensitivity, or light sensitivity.",
            "• High Contrast Theme (7:1+ contrast ratio)\n"
            "• Scalable typography up to 24px (extra large)\n"
            "• Prominent keyboard focus outlines (3px indigo)\n"
            "• Text-to-Speech (TTS) screen audio assistance\n"
            "• Automated image Alt-Text enforcement",
            "• 1.4.3 Contrast (Minimum) (AA)\n"
            "• 1.4.6 Contrast (Enhanced) (AAA)\n"
            "• 1.4.4 Resize Text (AA)\n"
            "• 2.1.1 Keyboard Navigation (A)"
        ),
        (
            "Easy Read Mode",
            "Learners with cognitive processing delays or language learners.",
            "• Simplified plain-language AI lesson summaries\n"
            "• Visual concept cards with dual-coding illustrations\n"
            "• Short bulleted sentences\n"
            "• Uncluttered UI mode (simplified navigation)",
            "• 3.1.5 Reading Level (AAA)\n"
            "• 3.1.2 Language of Parts (AA)\n"
            "• 1.3.1 Info and Relationships (A)"
        )
    ]

    col_widths = [Inches(1.2), Inches(1.5), Inches(2.3), Inches(1.5)]
    for row_idx, r_data in enumerate(preset_rows):
        row = table.add_row()
        for c_idx, text in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = COLOR_TEXT

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 3: EDUCATOR ACCESSIBILITY COMPLIANCE AUDIT ENGINE
    # ──────────────────────────────────────────────────────────────────────────
    add_h1("3. Educator Accessibility Audit Engine & Quality Checker")
    doc.add_paragraph(
        "To ensure all authored content meets accessibility standards before publication, ACESS incorporates an "
        "automated, real-time accessibility auditing engine (`accessibility-audit.ts`)."
    )

    add_h2("3.1 How the Audit Engine Works")
    doc.add_paragraph(
        "Whenever an educator writes, modifies, or saves a lesson in the Course Workspace, the audit engine runs "
        "a programmatic scan against the rendered HTML markup across six core evaluation axes:"
    )

    axes = [
        ("1. Heading Structure & Semantic Hierarchy: ", "Verifies that headings follow a logical sequence (H1 → H2 → H3) without skipping levels (e.g. H1 directly to H3), ensuring screen reader users can construct an accurate mental outline."),
        ("2. Non-Text Content & Image Alt-Text: ", "Scans all `<img>` tags for missing, empty, or generic alt attributes (e.g., 'image.png'). Images without descriptions are penalized."),
        ("3. Color Contrast & Legibility: ", "Evaluates text color values against background colors to ensure compliance with the 4.5:1 minimum contrast threshold for standard body text."),
        ("4. Cognitive Chunking & Paragraph Density: ", "Flags dense 'walls of text' (>150 words per paragraph) and recommends bulleted lists or micro-learning checkpoints to aid ADHD and Dyslexic comprehension."),
        ("5. Hyperlink Meaningfulness: ", "Flags ambiguous link anchor text (e.g., 'click here', 'read more') and requires descriptive link destinations (e.g., 'Read the WCAG Contrast Guidelines')."),
        ("6. Media Multimodality: ", "Checks whether video and audio assets include captions, transcripts, or supplementary interactive checkpoint questions.")
    ]
    for b_title, b_desc in axes:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(b_title)
        r1.bold = True
        r2 = bp.add_run(b_desc)

    add_h2("3.2 Scoring Algorithm & Compliance Bands")
    doc.add_paragraph("The audit engine computes a weighted score from 0 to 100 points, categorized into three distinct compliance bands:")

    tbl_score = doc.add_table(rows=4, cols=3)
    tbl_score.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_score.autofit = False
    s_headers = ["Score Band", "Status Category", "System Behavior & Guidance"]
    for i, h in enumerate(s_headers):
        c = tbl_score.cell(0, i)
        c.text = h
        set_cell_background(c, "1E1B4B")
        set_cell_margins(c, 100, 100, 100, 100)
        p = c.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    score_data = [
        ("80 – 100", "🟢 Good (Compliant)", "Lesson is fully accessible. Ready for immediate publishing to learners."),
        ("50 – 79", "🟡 Warning (Needs Attention)", "Minor issues detected (e.g., long paragraphs, suboptimal heading levels). Educator receives actionable recommendations."),
        ("0 – 49", "🔴 Critical (Non-Compliant)", "Severe accessibility violations (e.g., missing alt text, poor contrast). Highlighting provides inline fix guides.")
    ]
    s_widths = [Inches(1.2), Inches(2.2), Inches(3.1)]
    for r_idx, (b1, b2, b3) in enumerate(score_data):
        row = tbl_score.rows[r_idx + 1]
        for c_idx, val in enumerate([b1, b2, b3]):
            cell = row.cells[c_idx]
            cell.width = s_widths[c_idx]
            cell.text = val
            set_cell_margins(cell, 100, 100, 100, 100)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F9FAFB")
            p = cell.paragraphs[0]
            for r in p.runs:
                r.font.size = Pt(9.5)

    add_h2("3.3 How the Checker Benefits Educators and Learners")
    doc.add_paragraph("• For Educators: Removes the guesswork from accessible content creation. Instead of requiring deep WCAG technical expertise, educators receive instant, non-intrusive feedback with step-by-step guidance on how to fix issues.")
    doc.add_paragraph("• For Students: Guarantees a consistent, barrier-free learning experience. Neurodivergent students never encounter unformatted walls of text, broken contrast, or inaccessible multimedia.")

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 4: TEST PLAN — PLAYWRIGHT & K6
    # ──────────────────────────────────────────────────────────────────────────
    add_h1("4. Quality Assurance & Automated Testing Plan")
    doc.add_paragraph(
        "To ensure robust stability, security, and scalability under high concurrency, the testing strategy "
        "incorporates Playwright for End-to-End Functional testing and k6 for Load/Stress testing. "
        "Both tools can execute locally against Docker or remotely against the live Vercel deployment."
    )

    add_h2("4.1 Tool 1: Playwright (Functional & End-to-End Testing)")
    doc.add_paragraph("• What it is: A modern browser automation framework developed by Microsoft that automates Chromium, Firefox, and WebKit.")
    doc.add_paragraph("• What it validates: Multi-role authentication (Admin, Educator, Learner), portal RBAC routing guards, course enrollment, interactive quiz completion, accessibility preset switching, and certificate generation.")
    doc.add_paragraph("• What to expect: Fast parallel test runs with visual screenshots, network inspection traces, and an interactive HTML report.")

    add_h3("Playwright Step-by-Step Execution Guide:")
    p_code1 = doc.add_paragraph()
    p_code1.paragraph_format.left_indent = Inches(0.25)
    r = p_code1.add_run(
        "# 1. Install Playwright and browser binaries\n"
        "npm install -D @playwright/test\n"
        "npx playwright install --with-deps chromium\n\n"
        "# 2. Run tests against Localhost\n"
        "$env:BASE_URL=\"http://localhost:3000\"\n"
        "npx playwright test\n\n"
        "# 3. Run tests against Live Vercel Production\n"
        "$env:BASE_URL=\"https://your-acess-app.vercel.app\"\n"
        "npx playwright test\n\n"
        "# 4. View interactive HTML Report & video recordings\n"
        "npx playwright show-report"
    )
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(17, 24, 39)

    add_h2("4.2 Tool 2: k6 by Grafana (Load & Performance Stress Testing)")
    doc.add_paragraph("• What it is: A high-performance, developer-centric load testing engine written in Go, scripted in JavaScript.")
    doc.add_paragraph("• What it validates: API response latencies, serverless execution limits, database connection pooling under high concurrency, and SLA threshold compliance (e.g. 95% of requests completing under 500ms).")
    doc.add_paragraph("• What to expect: Live console metrics displaying Requests Per Second (RPS), p90/p95/p99 latency percentiles, error rates, and automated threshold pass/fail assertions.")

    add_h3("k6 Step-by-Step Execution Guide:")
    p_code2 = doc.add_paragraph()
    p_code2.paragraph_format.left_indent = Inches(0.25)
    r2 = p_code2.add_run(
        "# 1. Install k6 on Windows (via winget or portable exe)\n"
        "winget install k6 --source winget\n\n"
        "# 2. Run load test against Localhost\n"
        "k6 run -e BASE_URL=\"http://localhost:3000\" load-test.js\n\n"
        "# 3. Run load test against Live Vercel Deployment\n"
        "k6 run -e BASE_URL=\"https://your-acess-app.vercel.app\" load-test.js"
    )
    r2.font.name = 'Consolas'
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(17, 24, 39)

    add_h2("4.3 Concurrency Architecture: 50 Database Users vs 500 Concurrent VUs")
    doc.add_paragraph(
        "A common point of clarification in load testing: Virtual Users (VUs) represent concurrent network traffic streams, "
        "not distinct database account rows. 500 concurrent virtual users can test the system through:"
    )
    doc.add_paragraph("1. Public / Anonymous Endpoints (0 accounts needed): Landing page (`/`), public course catalog, and certificate verification (`/verify/[code]`) simulate hundreds of web visitors simultaneously.")
    doc.add_paragraph("2. Session Distribution across Seed Accounts: In k6, 500 virtual threads can round-robin across the 50 seeded database accounts (`VU % 50`), simulating multiple concurrent active sessions per user account.")

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 5: DATABASE STATE & SEEDED USER DIRECTORY
    # ──────────────────────────────────────────────────────────────────────────
    add_h1("5. Database Summary & Seeded User Directory")
    doc.add_paragraph(
        "The ACESS database is initialized with verified demo personas and realistic historical learning progress. "
        "All demo accounts use the standard testing password: **`AcessDemo#2026`**."
    )

    add_h2("5.1 Seeded User Directory & Scenario Roles")

    tbl_users = doc.add_table(rows=1, cols=5)
    tbl_users.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_users.autofit = False

    u_headers = ["Full Name", "Email Address", "System Role", "Preset / Need", "Assigned Test Scenario"]
    for i, h in enumerate(u_headers):
        c = tbl_users.cell(0, i)
        c.text = h
        set_cell_background(c, "1E1B4B")
        set_cell_margins(c, 100, 100, 80, 80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(255, 255, 255)

    user_rows = [
        ("Aliff Affandi", "aliff.admin@acess.edu.my", "admin", "None", "Super Admin; Full platform & user governance"),
        ("Nurul Izzah", "nurul.admin@acess.edu.my", "admin", "None", "Secondary Admin; Analytics & reporting oversight"),
        ("Rajesh Kumar", "rajesh.admin@acess.edu.my", "admin", "None", "Compliance Admin; Certificate & vetting auditor"),
        ("Dr. Siti Aminah", "siti.educator@acess.edu.my", "educator", "None", "Senior Educator; Author of Foundations & Reading"),
        ("Marcus Tan", "marcus.educator@acess.edu.my", "educator", "None", "STEM Educator; Author of ADHD Focus & Math"),
        ("Farah Nabilah", "farah.educator@acess.edu.my", "educator", "None", "Special Education Specialist; Author of Autism routines"),
        ("Amir Hakim", "amir.learner@acess.edu.my", "learner", "adhd", "Advanced Learner; 5 courses, 2 completed, 2 certs (100% done)"),
        ("Chong Mei Ling", "mei.learner@acess.edu.my", "learner", "dyslexia", "Mid-Progress Learner; 4 enrolled, 8/12 lessons completed"),
        ("Haziq Danial", "haziq.learner@acess.edu.my", "learner", "none", "New Learner; 2 enrolled, early onboarding stage"),
        ("Aisyah Nabila", "aisyah.learner@acess.edu.my", "learner", "dyslexia", "High Achiever; 4 enrolled, 3 completed, 5 achievements"),
        ("Priya Devi", "priya.learner@acess.edu.my", "learner", "autism", "Autism Routine Learner; 100% completed Autism courses"),
        ("Daniel Lim", "daniel.learner@acess.edu.my", "learner", "adhd", "At-Risk / Inactive Learner; 1 dropped course, low quiz scores")
    ]

    u_widths = [Inches(1.1), Inches(1.8), Inches(0.8), Inches(0.9), Inches(1.9)]
    for r_idx, u_data in enumerate(user_rows):
        row = tbl_users.add_row()
        for c_idx, val in enumerate(u_data):
            cell = row.cells[c_idx]
            cell.width = u_widths[c_idx]
            cell.text = val
            set_cell_margins(cell, 80, 80, 80, 80)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(8.5)
                r.font.color.rgb = COLOR_TEXT

    add_h2("5.2 Summary of Initialized Database Metrics")
    doc.add_paragraph("The seeded database comprises complete end-to-end telemetry reflecting authentic platform usage:")

    tbl_meta = doc.add_table(rows=1, cols=4)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_meta.autofit = False

    m_headers = ["Category", "Metric Count", "Category", "Metric Count"]
    for i, h in enumerate(m_headers):
        c = tbl_meta.cell(0, i)
        c.text = h
        set_cell_background(c, "4F46E5")
        set_cell_margins(c, 100, 100, 100, 100)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    meta_rows = [
        ("Total Database Tables", "42 Public Tables", "Enrolled Student Tracks", "21 Enrollments"),
        ("Admin Accounts", "3 Accounts", "Lesson Progress Records", "64 Progress Rows"),
        ("Educator Accounts", "3 Accounts", "Quiz Submissions & Attempts", "19 Quiz Attempts"),
        ("Learner Accounts", "6 Personas", "Quiz Question Answer Records", "60 Question Answers"),
        ("Published Courses", "8 Active Courses", "Defined Platform Achievements", "27 Achievements"),
        ("Draft / Review Courses", "3 Courses", "Earned Student Achievements", "46 Awarded Badges"),
        ("Curriculum Chapters", "14 Chapters", "Verified Certificates Issued", "7 Credentials (1 Revoked)"),
        ("Published Lessons", "29 Lessons", "Adaptive Interactions Tracked", "34 Telemetry Events"),
        ("Interactive Activities", "10 Checkpoints", "System Notifications Dispatched", "104 Notifications"),
        ("H5P & Media Embeds", "4 Embeds", "Accessibility Good Band (>80)", "22 Lessons (82-100 score)")
    ]

    m_widths = [Inches(1.8), Inches(1.4), Inches(1.9), Inches(1.4)]
    for r_idx, m_data in enumerate(meta_rows):
        row = tbl_meta.add_row()
        for c_idx, val in enumerate(m_data):
            cell = row.cells[c_idx]
            cell.width = m_widths[c_idx]
            cell.text = val
            set_cell_margins(cell, 80, 80, 80, 80)
            if r_idx % 2 == 1:
                set_cell_background(cell, "F9FAFB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = COLOR_TEXT

    # Output path
    output_path = os.path.join(os.getcwd(), "ACESS_Testing_and_Improvement_Report.docx")
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")

if __name__ == "__main__":
    create_document()
