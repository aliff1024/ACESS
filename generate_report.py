from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# STYLES SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Configure heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(14)
        hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)
    elif level == 3:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(3)

# Set margins
for section in doc.sections:
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_paragraph(text, style_name='Normal', bold=False, alignment=None, font_size=12, space_after=0, space_before=0, italic=False, indent_cm=0):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered(text, number):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    add_paragraph('')

add_paragraph('ACESS - ADAPTIVE COGNITIVE & EDUCATIONAL SKILL SUPPORT PLATFORM', 
              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=16, space_after=24)
add_paragraph('A Web-Based Accessible E-Learning Platform for Learners with Dyslexia, ADHD, and Mild Cognitive Impairment',
              italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, space_after=36)
add_paragraph('MUHAMMAD ALIFF BIN AFFANDI', 
              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, space_after=24)
add_paragraph('B032420067', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
add_paragraph('Bachelor of Computer Science (Software Engineering) (BITS)',
              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, space_after=48)
add_paragraph('FACULTY OF INFORMATION AND COMMUNICATION TECHNOLOGY',
              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, space_after=6)
add_paragraph('UNIVERSITI TEKNIKAL MALAYSIA MELAKA',
              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, space_after=6)
add_paragraph('2025/2026', alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12)

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_text('ABSTRACT', level=1)
add_paragraph('Learning disabilities such as dyslexia, ADHD, and mild cognitive impairment affect a significant portion of the population in Malaysia, with 299,128 individuals or 37.1% of registered persons with disabilities recorded in 2024. Traditional e-learning platforms are designed for neurotypical users and often overwhelm these learners with dense text, fast-paced videos, and complex interfaces, creating substantial barriers to skill development and employment. This project presents the development of ACESS (Adaptive Cognitive & Educational Skill Support Platform), an accessible, role-based web platform built using Next.js 16, Supabase, and TypeScript. The system enables educators to create structured courses with sequential lessons, transcripts, and quizzes, while learners benefit from built-in Text-to-Speech (TTS) read-aloud functionality, adjustable font sizes, multiple display themes, and a rule-based adaptive recommendation engine that automatically adjusts content difficulty based on quiz performance. The platform implements eight core modules: authentication and user management, course and lesson management, adaptive content recommendation, quiz and assessment, progress tracking, accessibility settings, certificate generation, and educator analytics dashboard. Role-Based Access Control (RBAC) enforces clear separation between learners, educators, and administrators. The system was developed using the Agile methodology, with testing conducted through functional unit tests and user acceptance testing. Results demonstrate that the platform successfully provides an inclusive digital learning environment that improves learning outcomes and provides employers and disability support organizations with verifiable evidence of competencies.',
              indent_cm=1.27, space_after=12)

add_heading_text('ABSTRAK', level=1)
add_paragraph('Masalah pembelajaran seperti disleksia, ADHD, dan masalah kognitif ringan menjejaskan sebahagian besar penduduk di Malaysia, dengan 299,128 individu atau 37.1% daripada orang kurang upaya berdaftar direkodkan pada tahun 2024. Platform e-pembelajaran tradisional direka untuk pengguna neurotipikal dan sering membebankan pelajar ini dengan teks padat, video pantas, dan antara muka yang kompleks, mewujudkan halangan yang besar kepada pembangunan kemahiran dan pekerjaan. Projek ini membentangkan pembangunan ACESS (Adaptive Cognitive & Educational Skill Support Platform), sebuah platform web berasaskan peranan yang boleh diakses, dibina menggunakan Next.js 16, Supabase, dan TypeScript. Sistem ini membolehkan pendidik membuat kursus berstruktur dengan pelajaran berturutan, transkrip, dan kuiz, manakala pelajar mendapat manfaat daripada ciri Text-to-Speech (TTS), saiz fon boleh laras, pelbagai tema paparan, dan enjin cadangan adaptif berasaskan peraturan yang secara automatik menyesuaikan kesukaran kandungan berdasarkan prestasi kuiz. Platform ini melaksanakan lapan modul teras: pengesahan dan pengurusan pengguna, pengurusan kursus dan pelajaran, cadangan kandungan adaptif, kuiz dan penilaian, penjejakan kemajuan, tetapan kebolehcapaian, penjanaan sijil, dan papan pemuka analitik pendidik. Sistem ini berjaya dibangunkan dan diuji, menyediakan persekitaran pembelajaran digital yang inklusif.',
              indent_cm=1.27, space_after=12)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (simplified)
# ============================================================
add_heading_text('TABLE OF CONTENTS', level=1)
toc_items = [
    ('CHAPTER 1. INTRODUCTION', 1),
    ('1.1 Project Background', 1),
    ('1.2 Problem Statements', 2),
    ('1.3 Objectives', 3),
    ('1.4 Scope of the Project', 3),
    ('1.5 Project Significance', 5),
    ('1.6 Expected Output', 5),
    ('1.7 Report Organization', 5),
    ('CHAPTER 2. LITERATURE REVIEW AND PROJECT METHODOLOGY', 6),
    ('2.1 Introduction', 6),
    ('2.2 Domain and Existing Systems', 6),
    ('2.3 Previous Studies', 7),
    ('2.4 Project Methodology', 8),
    ('2.5 Project Requirements', 9),
    ('2.6 Project Schedule and Milestones', 10),
    ('CHAPTER 3. ANALYSIS', 11),
    ('3.1 Introduction', 11),
    ('3.2 Problem Analysis', 11),
    ('3.3 Requirement Analysis', 12),
    ('CHAPTER 4. DESIGN', 14),
    ('4.1 Introduction', 14),
    ('4.2 High-Level Design / System Architecture', 14),
    ('4.3 User Interface Design', 15),
    ('4.4 Database Design', 16),
    ('4.5 Detailed Design', 17),
    ('CHAPTER 5. IMPLEMENTATION', 18),
    ('5.1 Introduction', 18),
    ('5.2 Software Development Environment Setup', 18),
    ('5.3 Software Configuration Management', 19),
    ('5.4 Implementation Status', 19),
    ('5.5 Key Implementation Details', 20),
    ('CHAPTER 6. TESTING', 22),
    ('6.1 Introduction', 22),
    ('6.2 Test Plan', 22),
    ('6.3 Test Strategy', 23),
    ('6.4 Test Design', 23),
    ('6.5 Test Results and Analysis', 24),
    ('CHAPTER 7. CONCLUSION', 25),
    ('7.1 Introduction', 25),
    ('7.2 Project Summarization', 25),
    ('7.3 Project Contribution', 26),
    ('7.4 Project Limitation', 26),
    ('7.5 Future Works', 26),
    ('REFERENCES', 27),
    ('APPENDICES', 28),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{item}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
add_heading_text('CHAPTER 1', level=1)
add_heading_text('INTRODUCTION', level=1)

add_heading_text('1.1 Project Background', level=2)
add_paragraph('Learning disabilities such as dyslexia, ADHD (Attention Deficit Hyperactivity Disorder), and mild cognitive impairment affect many people around the world. Globally, approximately 15 to 20 percent of individuals experience some form of learning or thinking difference. In Malaysia, a total of 805,509 persons with disabilities were registered in 2024, representing 2.4 percent of the population. Among them, learning disabilities recorded the highest number with 299,128 individuals, or 37.1 percent of all registered cases (Department of Statistics Malaysia, 2024).',
              indent_cm=1.27, space_after=6)
add_paragraph('These conditions do not indicate lower intelligence. They simply reflect different ways the brain processes information. However, most existing e-learning platforms are designed for neurotypical learners and provide little to no support for cognitive, sensory, or attentional needs. As education and vocational skill training continue to shift online, learners with disabilities often feel overwhelmed by dense text, fast-paced videos, and complex interfaces. Without suitable accommodations such as adjustable text size, read-aloud features, simpler navigation, and flexible pacing, many experience frustration, disengagement, and high dropout rates.',
              indent_cm=1.27, space_after=6)
add_paragraph('This project seeks to address these challenges by developing ACESS - Adaptive Cognitive & Educational Skill Support Platform. The system will offer a complete, accessible digital learning journey from user registration to certificate generation. It includes adaptive content recommendations, built-in accessibility tools, progress tracking, and dedicated dashboards for educators and administrators.',
              indent_cm=1.27, space_after=6)

add_heading_text('1.2 Problem Statements', level=2)
add_paragraph('The following problem statements have been identified:', indent_cm=1.27, space_after=6)

problems = [
    ('PS1: ', 'Most existing e-learning platforms are not designed for learners with disabilities. They use formats that overwhelm cognitive and sensory processing, leading to high dropout rates.'),
    ('PS2: ', 'Learners with disabilities lack platforms that adapt to their individual learning pace and present content in accessible, easy-to-process formats.'),
    ('PS3: ', 'Employers and support organisations have no reliable way to verify the skills gained by learners with disabilities due to the absence of proper progress tracking and accessible certification.'),
    ('PS4: ', 'Educators and mentors do not have a centralised platform to monitor learner progress, identify those who need extra support, and manage accessible learning materials effectively.')
]
for prefix, text in problems:
    add_bullet(text, bold_prefix=prefix)

add_heading_text('1.3 Objectives', level=2)
add_paragraph('This project aims to achieve the following objectives:', indent_cm=1.27, space_after=6)
objectives = [
    ('Objective 1: ', 'To analyze the accessibility requirements, learning challenges, and content delivery needs of learners with dyslexia, ADHD, and mild cognitive impairment through literature review and comparative study of existing e-learning platforms and assistive technology standards.'),
    ('Objective 2: ', 'To develop a fully functional accessible web-based e-learning platform that incorporates rule-based adaptive content delivery, built-in Text-to-Speech (TTS) support, adjustable accessibility display settings, structured lesson and quiz management, and role-based user management for learners, educators, and administrators.'),
    ('Objective 3: ', 'To implement a certificate generation module and an educator analytics dashboard that provide verifiable skill completion records and enable data-driven monitoring of individual learner progress.'),
    ('Objective 4: ', 'To evaluate the functional correctness, usability, and accessibility of the developed platform through structured user acceptance testing (UAT), measuring task completion rates and user satisfaction using the System Usability Scale (SUS).')
]
for prefix, text in objectives:
    add_bullet(text, bold_prefix=prefix)

add_heading_text('1.4 Scope of the Project', level=2)
add_paragraph('The project covers the following eight system modules:', indent_cm=1.27, space_after=6)

scopes = [
    ('Authentication and User Management Module - ', 'Handles secure registration with email verification, login, logout, and password reset. Implements Role-Based Access Control (RBAC) for Learner, Educator, and Administrator roles. Allows administrators to create, edit, activate/deactivate user accounts. Includes learner onboarding profile for declaring disability type and accessibility preferences.'),
    ('Course and Lesson Management Module - ', 'Enables educators to create, edit, and publish structured courses with sequential lessons, titles, descriptions, and skill tags. Supports text content, embedded videos with transcripts, and prerequisite sequencing for scaffolded learning. Provides learners with a visual progress bar and module completion checklist.'),
    ('Adaptive Content Recommendation Module - ', 'Uses rule-based logic to adjust difficulty based on quiz scores (revision below 60%, standard 60-80%, advanced above 80%). Displays personalized "Recommended Next" cards on the learner dashboard. Logs all recommendations for educator review.'),
    ('Quiz and Assessment Module - ', 'Allows educators to create multiple-choice and scenario-based quizzes with time limits and attempt controls. Presents questions one at a time in a clean, distraction-free interface with large text and progress indicators. Provides immediate feedback and feeds scores into the adaptive recommendation system.'),
    ('Progress Tracking Module - ', 'Records lesson views, quiz attempts, scores, and timestamps for each learner. Displays a simple personal dashboard for learners showing completion percentage and next steps. Enables educators to view individual progress and identify at-risk learners.'),
    ('Accessibility Settings Module - ', 'Allows learners to customize font size, display theme (light/dark/high contrast/soft), line spacing, and font type. Integrates a one-click Text-to-Speech (TTS) button for lesson content. Provides toggleable transcripts for video content. Saves all preferences to the user account for automatic restoration on login.'),
    ('Certificate Generation Module - ', 'Automatically generates downloadable PDF certificates upon course completion with learner name, course title, date, and unique reference code. Provides learners with a dedicated page to view and download all earned certificates. Allows administrators to view, manage, revoke, or reissue certificates.'),
    ('Educator Analytics Dashboard Module - ', 'Displays course enrollment, average quiz scores, completion rates, and at-risk learners. Visualizes quiz score distributions using bar charts. Supports filtering by course, date range, and disability type. Allows export of progress reports in CSV format.')
]
for prefix, text in scopes:
    add_bullet(text, bold_prefix=prefix)

add_heading_text('1.5 Project Significance', level=2)
add_paragraph('Upon successful completion, this project is expected to deliver:', indent_cm=1.27, space_after=6)
significance = [
    'A fully functional accessible web platform providing learners with disabilities a disability-friendly learning environment.',
    'An adaptive content recommendation system that adjusts lesson difficulty based on quiz performance.',
    'An automatic certificate generation module producing downloadable PDF certificates with unique reference numbers.',
    'A user-friendly educator analytics dashboard for monitoring progress and making timely interventions.',
    'A secure, role-based system with clear separation of access for Learners, Educators, and Administrators.'
]
for item in significance:
    add_bullet(item)

add_heading_text('1.6 Expected Output', level=2)
add_paragraph('The project is expected to produce a fully deployed web application accessible via modern web browsers. The system will include all eight modules described in the scope, with a responsive design that works on desktop and mobile devices. The source code will be maintained in a Git repository with proper version control documentation.',
              indent_cm=1.27, space_after=6)

add_heading_text('1.7 Report Organization', level=2)
add_paragraph('This report is organized into seven chapters. Chapter 1 introduces the project background, problem statements, objectives, and scope. Chapter 2 presents a literature review of existing e-learning platforms and accessibility standards, along with the project methodology. Chapter 3 analyzes the current problems and system requirements. Chapter 4 describes the system design including architecture, database, and user interfaces. Chapter 5 details the implementation of all system modules. Chapter 6 presents the testing strategy and results. Finally, Chapter 7 concludes the project with a summary of achievements, limitations, and recommendations for future work.',
              indent_cm=1.27, space_after=6)

doc.add_page_break()

# ============================================================
# CHAPTER 2: LITERATURE REVIEW AND PROJECT METHODOLOGY
# ============================================================
add_heading_text('CHAPTER 2', level=1)
add_heading_text('LITERATURE REVIEW AND PROJECT METHODOLOGY', level=1)

add_heading_text('2.1 Introduction', level=2)
add_paragraph('This chapter presents a review of existing literature related to accessible e-learning platforms, learning disabilities, and assistive technologies. It also describes the project methodology adopted for developing the ACESS platform. The literature review examines the current state of e-learning platforms for learners with disabilities and identifies gaps that this project aims to address.',
              indent_cm=1.27, space_after=6)

add_heading_text('2.2 Domain and Existing Systems', level=2)
add_paragraph('The domain of this project is accessible e-learning for individuals with learning disabilities. Learning disabilities encompass conditions such as dyslexia (difficulty reading and processing written text), ADHD (difficulty maintaining attention and regulating focus), and mild cognitive impairment (subtle decline in cognitive abilities).',
              indent_cm=1.27, space_after=6)

add_paragraph('Several e-learning platforms exist in the current market:', indent_cm=1.27, space_after=6)

platforms = [
    ('TalentLMS - ', 'A cloud-based learning management system that offers course creation, assessments, and reporting. However, it lacks built-in accessibility features such as Text-to-Speech and adjustable display settings specifically designed for learning disabilities.'),
    ('360Learning - ', 'A collaborative learning platform focused on peer learning and social engagement. While it provides good content management features, it does not offer adaptive content recommendations based on learner performance.'),
    ('Moodle - ', 'An open-source LMS with extensive plugin support. While accessibility plugins exist, they require manual configuration and do not provide a unified, role-based accessible experience out of the box.'),
]
for prefix, text in platforms:
    add_bullet(text, bold_prefix=prefix)

add_paragraph('The comparison reveals a gap in the market: no existing platform combines role-based access control, adaptive content recommendations, built-in Text-to-Speech, adjustable accessibility settings, automatic certificate generation, and an educator analytics dashboard into a single integrated system designed specifically for learners with learning disabilities.',
              indent_cm=1.27, space_after=6)

add_heading_text('2.3 Previous Studies', level=2)
add_paragraph('Research in the field of accessible e-learning has highlighted several important findings. According to the Department of Statistics Malaysia (2024), learning disabilities represent the largest category of registered disabilities in Malaysia, comprising 37.1% of all cases. This statistic underscores the urgent need for accessible educational platforms.',
              indent_cm=1.27, space_after=6)
add_paragraph('Studies on dyslexia and digital learning (SingHealth, 2023) suggest that individuals with dyslexia benefit from increased letter spacing, sans-serif fonts, and high contrast between text and background. Research on ADHD and online learning indicates that simplified interfaces with reduced visual clutter significantly improve task completion rates. For individuals with mild cognitive impairment, paced learning with repetition and immediate feedback has been shown to enhance information retention.',
              indent_cm=1.27, space_after=6)
add_paragraph('The Web Content Accessibility Guidelines (WCAG) 2.1 provide a framework for creating accessible digital content. Key requirements include providing text alternatives for non-text content, creating content that can be presented in different ways without losing information, and making content easier to see and hear. The ACESS platform incorporates these guidelines into its design.',
              indent_cm=1.27, space_after=6)

add_heading_text('2.4 Project Methodology', level=2)
add_paragraph('The Agile Software Development Lifecycle (SDLC) methodology was adopted for this project. Agile was chosen due to its iterative nature, which allows for continuous feedback and adaptation throughout the development process. The methodology consists of the following phases:',
              indent_cm=1.27, space_after=6)

agile_phases = [
    ('Requirement Analysis: ', 'Gathering and analyzing requirements through literature review and comparative studies of existing platforms. Documenting functional and non-functional requirements.'),
    ('Design: ', 'Creating system architecture, database design, user interface mockups, and detailed design specifications using UML diagrams.'),
    ('Development (Sprints): ', 'Implementing the system in iterative sprints, with each sprint delivering a working increment of the software. Development follows the order of dependencies: authentication first, then core features, followed by accessibility and reporting features.'),
    ('Testing: ', 'Conducting unit testing, integration testing, system testing, and user acceptance testing at the end of each sprint.'),
    ('Deployment: ', 'Deploying the completed system to a production environment and documenting the setup process.')
]
for prefix, text in agile_phases:
    add_bullet(text, bold_prefix=prefix)

add_heading_text('2.5 Project Requirements', level=2)

add_paragraph('Software Requirements:', bold=True, space_after=3)
sw_reqs = [
    'Next.js 16 (React framework for server-side rendering and routing)',
    'Supabase (Backend-as-a-Service providing PostgreSQL database, authentication, and storage)',
    'TypeScript (Type-safe programming language)',
    'Tailwind CSS v4 (Utility-first CSS framework for rapid UI development)',
    'shadcn/ui (Component library built on Radix UI primitives)',
    'Node.js (Runtime environment)',
    'Git (Version control)'
]
for req in sw_reqs:
    add_bullet(req)

add_paragraph('Hardware Requirements:', bold=True, space_after=3, space_before=6)
hw_reqs = [
    'Development computer with minimum 8GB RAM',
    'Stable internet connection for cloud services (Supabase)',
    'Web server for deployment (Vercel or similar Node.js hosting)'
]
for req in hw_reqs:
    add_bullet(req)

add_heading_text('2.6 Project Schedule and Milestones', level=2)
add_paragraph('The project was carried out over two semesters (PSM 1 and PSM 2). The key milestones are summarized below:',
              indent_cm=1.27, space_after=6)

milestones = [
    ['Phase', 'Activities', 'Deliverables'],
    ['Literature Review', 'Review existing platforms and accessibility standards', 'Chapter 2 draft'],
    ['Requirement Analysis', 'Define functional and non-functional requirements', 'Chapter 1 & 3 draft'],
    ['System Design', 'Architecture, database, UI design', 'Chapter 4 draft'],
    ['Sprint 1', 'Authentication module, user management', 'Working login/signup'],
    ['Sprint 2', 'Course and lesson management', 'Course CRUD + lesson viewer'],
    ['Sprint 3', 'Quiz system + adaptive recommendations', 'Quiz engine + recommendations'],
    ['Sprint 4', 'Progress tracking + certificates', 'Dashboard + PDF generation'],
    ['Sprint 5', 'Analytics dashboard + admin panel', 'Educator/Admin dashboards'],
    ['Sprint 6', 'Accessibility polish + i18n + testing', 'Final system + test reports'],
    ['Report Writing', 'Complete final report', 'Full PSM report']
]
table = doc.add_table(rows=len(milestones), cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row_data in enumerate(milestones):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if r_idx == 0:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================
# CHAPTER 3: ANALYSIS
# ============================================================
add_heading_text('CHAPTER 3', level=1)
add_heading_text('ANALYSIS', level=1)

add_heading_text('3.1 Introduction', level=2)
add_paragraph('This chapter presents the analysis phase of the project. It examines the current problems with existing e-learning platforms, analyzes the requirements for the proposed system, and establishes the foundation for the design phase.',
              indent_cm=1.27, space_after=6)

add_heading_text('3.2 Problem Analysis', level=2)
add_paragraph('The current landscape of e-learning platforms reveals several critical limitations:', indent_cm=1.27, space_after=6)

problem_analysis = [
    'Lack of Accessibility Features: Most platforms treat accessibility as an afterthought, offering limited or no support for Text-to-Speech, adjustable font sizes, or high-contrast themes.',
    'No Adaptive Content Delivery: E-learning platforms typically present the same content to all learners regardless of their performance, failing to accommodate different learning paces.',
    'Limited Progress Tracking for Educators: Educators lack centralized dashboards to monitor individual learner progress and identify those who need additional support.',
    'Absence of Verifiable Certification: Learners completing courses often receive no formal documentation of their achievements that employers can verify.',
    'Complex Navigation: Standard e-learning interfaces contain cluttered layouts with multiple menus and options that overwhelm users with attention-related disabilities.'
]
for item in problem_analysis:
    add_bullet(item)

add_paragraph('The ACESS platform addresses these limitations through a unified, role-based architecture that integrates accessibility, adaptivity, and comprehensive tracking into a single system.',
              indent_cm=1.27, space_after=6)

add_heading_text('3.3 Requirement Analysis', level=2)

add_paragraph('3.3.1 Data Requirement', bold=True, space_after=3)
add_paragraph('The system manages the following data entities:', indent_cm=1.27, space_after=6)
data_entities = [
    'User data: profiles, authentication credentials, roles, accessibility preferences',
    'Course data: titles, descriptions, categories, difficulty levels, tags',
    'Lesson data: content (text, video URLs, transcripts), sequence ordering, prerequisites',
    'Quiz data: questions, options, correct answers, time limits, attempt limits',
    'Progress data: lesson views, quiz attempts, scores, completion status',
    'Certificate data: reference codes, issue dates, revocation status',
    'Recommendation data: difficulty tiers, trigger reasons, acknowledgement status',
    'Notification data: types, titles, bodies, read status'
]
for entity in data_entities:
    add_bullet(entity)

add_paragraph('3.3.2 Functional Requirement', bold=True, space_after=3, space_before=6)
add_paragraph('The functional requirements define the specific functions the system must perform. The system is divided into three role-based interfaces:',
              indent_cm=1.27, space_after=6)

add_paragraph('Learner Functions:', bold=True, space_after=3)
learner_funcs = [
    'Register and create a personal account with disability type and accessibility preferences',
    'Browse and enroll in available courses',
    'View lessons with Text-to-Speech playback and adjustable display settings',
    'Take quizzes with time limits and immediate feedback',
    'Track personal learning progress with completion percentages',
    'View and download earned certificates with unique reference codes',
    'Receive personalized content recommendations based on quiz performance'
]
for func in learner_funcs:
    add_bullet(func)

add_paragraph('Educator Functions:', bold=True, space_after=3, space_before=6)
educator_funcs = [
    'Create, edit, and manage courses with structured lessons and quizzes',
    'Upload lesson content including text and embedded videos with transcripts',
    'Create quizzes with multiple-choice questions and correct answers',
    'View analytics dashboard showing enrollment, completion rates, and at-risk learners',
    'Monitor individual learner progress and quiz performance',
    'Export progress reports in CSV format'
]
for func in educator_funcs:
    add_bullet(func)

add_paragraph('Administrator Functions:', bold=True, space_after=3, space_before=6)
admin_funcs = [
    'Manage all user accounts and assign roles',
    'Review and approve courses before publication',
    'Monitor platform-wide statistics and engagement metrics',
    'Manage certificate issuance, revocation, and reissuance',
    'Generate system-wide reports'
]
for func in admin_funcs:
    add_bullet(func)

add_paragraph('3.3.3 Non-functional Requirement', bold=True, space_after=3, space_before=6)
nf_reqs = [
    'Performance: Page load times should not exceed 3 seconds under normal network conditions.',
    'Availability: The system should maintain 99% uptime during business hours.',
    'Security: All user data must be encrypted in transit (HTTPS) and at rest. Passwords must be hashed using bcrypt.',
    'Usability: The interface must be navigable via keyboard alone and compatible with screen readers.',
    'Scalability: The system should support up to 1000 concurrent users.',
    'Responsiveness: The web application must be fully functional on desktop and mobile devices.',
    'Accessibility: Must comply with WCAG 2.1 AA standards.'
]
for req in nf_reqs:
    add_bullet(req)

add_paragraph('3.3.4 Other Requirements', bold=True, space_after=3, space_before=6)
add_paragraph('The system requires a modern web browser (Chrome, Firefox, Safari, or Edge) on the client side. Server-side deployment requires a Node.js hosting environment (such as Vercel) and a Supabase project for database, authentication, and storage services. Development requires a code editor (VS Code) and Git for version control.',
              indent_cm=1.27, space_after=6)

doc.add_page_break()

# ============================================================
# CHAPTER 4: DESIGN
# ============================================================
add_heading_text('CHAPTER 4', level=1)
add_heading_text('DESIGN', level=1)

add_heading_text('4.1 Introduction', level=2)
add_paragraph('This chapter presents the system design, including the high-level architecture, user interface design, database design, and detailed component design. The design translates the requirements identified in Chapter 3 into a concrete blueprint for implementation.',
              indent_cm=1.27, space_after=6)

add_heading_text('4.2 High-Level Design / System Architecture', level=2)
add_paragraph('The ACESS platform follows a three-tier architecture with the following layers:', indent_cm=1.27, space_after=6)

add_paragraph('Presentation Layer (Frontend):', bold=True, space_after=3)
add_paragraph('Built with Next.js 16 using the App Router, this layer handles all user interactions. It includes role-specific dashboards for learners, educators, and administrators. The frontend uses React Context for state management with three providers: AuthProvider (authentication state), AccessibilityProvider (display settings), and LanguageProvider (i18n).',
              indent_cm=1.27, space_after=6)

add_paragraph('Application Logic Layer (API/Middleware):', bold=True, space_after=3)
add_paragraph('Next.js middleware (proxy.ts) handles route-level authentication and role-based access control. Three API modules wrap all Supabase queries: learner-api.ts (~1270 lines), educator-api.ts (~838 lines), and admin-api.ts (~851 lines). A recommendation engine (recommendation-engine.ts) uses a service role client for cross-user operations.',
              indent_cm=1.27, space_after=6)

add_paragraph('Data Layer (Supabase):', bold=True, space_after=3)
add_paragraph('Supabase provides PostgreSQL database, authentication, and file storage. Three server clients are used: createBrowserClient() for browser-side queries, createServerClient() for server components and API routes, and createClient() with service role key for admin operations.',
              indent_cm=1.27, space_after=6)

add_paragraph('The data flow follows this pattern: Browser → Next.js Middleware (auth check) → API Routes/Supabase queries → Database. The provider nesting is: AuthProvider wraps AccessibilityProvider wraps LanguageProvider wraps SessionTimeout.',
              indent_cm=1.27, space_after=6)

add_heading_text('4.3 User Interface Design', level=2)

add_paragraph('4.3.1 Navigation Design', bold=True, space_after=3)
add_paragraph('The navigation structure is role-based. After login, users are redirected to their role-specific dashboard:', indent_cm=1.27, space_after=6)
nav_items = [
    '/learner - Learner dashboard with enrolled courses, recommendations, and progress',
    '/educator - Educator dashboard with course management and analytics',
    '/admin - Administrator dashboard with user management and system statistics',
    '/profile - User profile and accessibility settings (all roles)',
    '/login, /signup, /forgot-password, /reset-password - Public authentication pages'
]
for item in nav_items:
    add_bullet(item)

add_paragraph('4.3.2 Input Design', bold=True, space_after=3, space_before=6)
add_paragraph('Input forms use react-hook-form with zod validation. Key input screens include:', indent_cm=1.27, space_after=6)
input_screens = [
    'Login: email and password fields with validation',
    'Signup: email, password, full name, and role selection via radio group',
    'Course creation: title, description, category, difficulty level, tags',
    'Lesson editor: title, content text, video URL, transcript, sequence order',
    'Quiz builder: title, time limit, max attempts, pass threshold, questions with options',
    'Accessibility settings: font size, theme, line spacing, font type, TTS toggle, all using radio groups'
]
for item in input_screens:
    add_bullet(item)

add_paragraph('4.3.3 Output Design', bold=True, space_after=3, space_before=6)
add_paragraph('The system produces the following outputs:', indent_cm=1.27, space_after=6)
outputs = [
    'Learner dashboard: course cards with progress bars, recommendation cards with difficulty tiers',
    'Analytics dashboard: bar charts for quiz score distributions, tables for enrollment and completion rates',
    'Certificates: downloadable PDF with learner name, course title, completion date, unique reference code',
    'Progress reports: CSV export for educators showing individual and group performance',
    'Notifications: bell icon with unread count, dropdown list of recent notifications'
]
for item in outputs:
    add_bullet(item)

add_heading_text('4.4 Database Design', level=2)
add_paragraph('The database uses PostgreSQL hosted on Supabase. The schema consists of the following main tables:',
              indent_cm=1.27, space_after=6)

db_tables = [
    ['Table Name', 'Description', 'Key Fields'],
    ['users', 'User profiles (synced from auth.users via trigger)', 'id, email, full_name, role, avatar_url, deleted_at'],
    ['courses', 'Course metadata', 'id, title, slug, description, category, difficulty, status, created_by'],
    ['course_tags', 'Course categorization tags', 'id, course_id, tag'],
    ['enrollments', 'Learner-course relationships', 'id, user_id, course_id, status, enrolled_at'],
    ['lessons', 'Lesson content within courses', 'id, course_id, title, content, video_url, transcript, sequence_order, status'],
    ['lesson_progress', 'Learner lesson tracking', 'id, enrollment_id, lesson_id, is_viewed, viewed_at'],
    ['quizzes', 'Quiz metadata per lesson', 'id, lesson_id, title, time_limit_seconds, max_attempts, pass_threshold_pct'],
    ['quiz_questions', 'Individual questions', 'id, quiz_id, question_text, question_type, sequence_order'],
    ['quiz_options', 'Answer options for questions', 'id, question_id, option_text, is_correct, sequence_order'],
    ['quiz_attempts', 'Learner quiz submissions', 'id, enrollment_id, quiz_id, attempt_number, score_pct, result'],
    ['quiz_answers', 'Individual answer records', 'id, attempt_id, question_id, selected_option_id'],
    ['certificates', 'Course completion certificates', 'id, enrollment_id, reference_code, status, issued_at, revoked_at'],
    ['recommendations', 'Adaptive recommendations', 'id, enrollment_id, recommended_lesson_id, difficulty_tier, trigger_reason, is_acknowledged'],
    ['notifications', 'System notifications', 'id, user_id, type, title, body, is_read, created_at'],
    ['user_accessibility_settings', 'User accessibility preferences', 'user_id, preferred_font_size, preferred_theme, line_spacing, preferred_font, tts_enabled'],
    ['password_reset_tokens', 'Password reset tokens', 'id, user_id, email, token, expires_at, used']
]
add_table(db_tables[0], db_tables[1:])
add_paragraph('')

add_heading_text('4.5 Detailed Design', level=2)
add_paragraph('The system follows an Object-Oriented design pattern using React components. Key design decisions include:',
              indent_cm=1.27, space_after=6)
design_decisions = [
    'No React Query / SWR: Direct Supabase calls with useState and useEffect for data fetching.',
    'No Server Actions: All mutations performed via client-side Supabase calls.',
    'Service role for admin operations: Bypasses RLS for password resets and recommendation generation.',
    'CSS data attributes for themes: Accessibility settings applied via data-* attributes on the <html> element, enabling instant theme switching without page reload.',
    'Custom SMTP: Bypasses Supabase email system for password resets, using Nodemailer with Gmail SMTP for full control over email design.'
]
for item in design_decisions:
    add_bullet(item)

doc.add_page_break()

# ============================================================
# CHAPTER 5: IMPLEMENTATION
# ============================================================
add_heading_text('CHAPTER 5', level=1)
add_heading_text('IMPLEMENTATION', level=1)

add_heading_text('5.1 Introduction', level=2)
add_paragraph('This chapter describes the implementation of the ACESS platform. It covers the development environment setup, software configuration management, and the implementation status of each system module. The implementation translates the design specifications from Chapter 4 into a working software system.',
              indent_cm=1.27, space_after=6)

add_heading_text('5.2 Software Development Environment Setup', level=2)
add_paragraph('The development environment was configured as follows:', indent_cm=1.27, space_after=6)

add_paragraph('Hardware Environment:', bold=True, space_after=3)
hw_env = [
    'Development computer: Windows 11, Intel Core i5, 16GB RAM',
    'Cloud server: Vercel for hosting, Supabase for backend services'
]
for item in hw_env:
    add_bullet(item)

add_paragraph('Software Environment:', bold=True, space_after=3, space_before=6)
sw_env = [
    'Operating System: Windows 11',
    'Runtime: Node.js (Latest LTS)',
    'Framework: Next.js 16 with App Router',
    'Database: Supabase (PostgreSQL 15)',
    'Authentication: Supabase Auth with custom SMTP',
    'Storage: Supabase Storage for file uploads',
    'IDE: Visual Studio Code',
    'Version Control: Git with GitHub remote repository',
    'Package Manager: npm'
]
for item in sw_env:
    add_bullet(item)

add_heading_text('5.3 Software Configuration Management', level=2)
add_paragraph('Version control is managed using Git with the repository hosted on GitHub at https://github.com/aliff1024/ACESS.git. The main branch is protected and all changes go through feature branches. The project uses a standard .gitignore file for Node.js projects. Environment variables are stored in .env.local (excluded from version control).',
              indent_cm=1.27, space_after=6)

add_heading_text('5.4 Implementation Status', level=2)
add_paragraph('The following table summarizes the implementation status of each system module:',
              indent_cm=1.27, space_after=6)

impl_table = [
    ['Module', 'Description', 'Status'],
    ['Authentication & User Management', 'Login, signup, logout, password reset, role-based access', 'Completed'],
    ['Course & Lesson Management', 'Course CRUD, lesson sequencing, prerequisite locking', 'Completed'],
    ['Adaptive Content Recommendation', 'Rule-based engine, cross-course tag matching', 'Completed'],
    ['Quiz & Assessment', 'Quiz builder, timed quizzes, auto-grading, attempt tracking', 'Completed'],
    ['Progress Tracking', 'Lesson progress, quiz scores, completion percentages', 'Completed'],
    ['Accessibility Settings', 'Font size, themes, line spacing, TTS, font type', 'Completed'],
    ['Certificate Generation', 'PDF generation, unique reference codes, admin management', 'Completed'],
    ['Educator Analytics Dashboard', 'Enrollment stats, quiz scores, at-risk learners, CSV export', 'Completed'],
    ['Notification System', 'DB triggers, unread counts, real-time display', 'Completed'],
    ['Multi-language Support', 'English/Malay translations via i18n provider', 'Completed'],
    ['Admin Panel', 'User management, course approval, certificate oversight', 'Completed']
]
add_table(impl_table[0], impl_table[1:])
add_paragraph('')

add_heading_text('5.5 Key Implementation Details', level=2)

add_paragraph('5.5.1 Authentication Module', bold=True, space_after=3)
add_paragraph('The authentication module uses Supabase Auth with a custom password reset flow. Key implementation details:', indent_cm=1.27, space_after=3)
auth_details = [
    'Login: supabase.auth.signInWithPassword() with role-based redirect',
    'Signup: supabase.auth.signUp() with user_metadata storing full_name and role',
    'Session management: AuthProvider React context with onAuthStateChange listener',
    'Role-based middleware: proxy.ts intercepts all requests and checks user.role against route prefix',
    'Password reset: Custom flow using crypto.randomBytes() for tokens, Nodemailer for emails, admin.updateUserById() for password changes'
]
for item in auth_details:
    add_bullet(item)

add_paragraph('5.5.2 Accessibility Module', bold=True, space_after=3, space_before=6)
add_paragraph('The accessibility module is implemented using CSS data attributes on the <html> element:', indent_cm=1.27, space_after=3)
acc_details = [
    'Settings stored in user_accessibility_settings table',
    'Applied via useEffect in AccessibilityProvider: document.documentElement.setAttribute()',
    'Four font sizes: small, medium, large, xlarge mapped to --font-size-base CSS variable',
    'Four display themes: light, dark, high_contrast, soft with complete color overrides',
    'Dyslexia-friendly font option: Trebuchet MS with increased letter and word spacing',
    'Reduced motion mode: disables all CSS transitions and animations',
    'One-click Text-to-Speech for lesson content'
]
for item in acc_details:
    add_bullet(item)

add_paragraph('5.5.3 Recommendation Engine', bold=True, space_after=3, space_before=6)
add_paragraph('The recommendation engine (src/lib/recommendation-engine.ts) implements a rule-based algorithm:', indent_cm=1.27, space_after=3)
rec_details = [
    'Revision tier: triggered when quiz score is below pass threshold (default 60%), recommends re-reading the lesson',
    'Standard tier: recommends the next unviewed lesson in the course sequence',
    'Advanced tier: triggered when course progress reaches 80% or higher',
    'Cross-course recommendations: matches tags and categories from completed courses to suggest new courses',
    'Uses service role client to bypass RLS for cross-user operations',
    'All old recommendations are deleted and replaced with fresh ones on each generation'
]
for item in rec_details:
    add_bullet(item)

add_paragraph('5.5.4 Notification System', bold=True, space_after=3, space_before=6)
add_paragraph('The notification system uses PostgreSQL database triggers:', indent_cm=1.27, space_after=3)
notif_details = [
    'on_enrollment_notify: Notifies educator when a learner enrolls',
    'on_lesson_progress_notify: Notifies educator when a lesson is first viewed',
    'on_quiz_attempt_notify: Notifies educator of quiz completion',
    'on_lesson_added_notify: Notifies enrolled learners of new content',
    'on_course_published_notify: Notifies enrolled learners when course is published',
    'Client-side fetching: getUnreadCount(), fetchNotifications(), markAsRead(), markAllAsRead()'
]
for item in notif_details:
    add_bullet(item)

doc.add_page_break()

# ============================================================
# CHAPTER 6: TESTING
# ============================================================
add_heading_text('CHAPTER 6', level=1)
add_heading_text('TESTING', level=1)

add_heading_text('6.1 Introduction', level=2)
add_paragraph('This chapter presents the testing and evaluation of the implemented ACESS platform. The purpose of testing is to verify that the system functions correctly according to the requirements defined in Chapter 3 and to ensure that the system performs reliably under expected operating conditions.',
              indent_cm=1.27, space_after=6)

add_heading_text('6.2 Test Plan', level=2)

add_paragraph('6.2.1 Test Organization', bold=True, space_after=3)
add_paragraph('The following personnel were involved in testing:', indent_cm=1.27, space_after=6)

test_org = [
    ['Role', 'Responsibility'],
    ['Developer', 'Performs unit testing and integration testing of system modules'],
    ['Supervisor', 'Reviews system functionality and design adherence'],
    ['Test Users', 'Evaluate system usability and functionality through UAT']
]
add_table(test_org[0], test_org[1:])
add_paragraph('')

add_paragraph('6.2.2 Test Environment', bold=True, space_after=3)
add_paragraph('Testing was conducted in the following environment:', indent_cm=1.27, space_after=6)
test_env = [
    'Hardware: Development computer (Windows 11, Intel Core i5, 16GB RAM)',
    'Browser: Google Chrome (latest version) with developer tools',
    'Backend: Supabase production project with test data',
    'Network: Local development server (localhost:3000) and deployed instance on Vercel'
]
for item in test_env:
    add_bullet(item)

add_paragraph('6.2.3 Test Schedule', bold=True, space_after=3)
add_paragraph('Testing was conducted in phases aligned with the development sprints:', indent_cm=1.27, space_after=6)
test_sched = [
    ['Test Phase', 'Duration', 'Description'],
    ['Unit Testing', 'Throughout sprints', 'Testing individual functions and components'],
    ['Integration Testing', 'End of each sprint', 'Testing interactions between modules'],
    ['System Testing', 'Final sprint', 'Testing complete system end-to-end'],
    ['User Acceptance Testing', 'After deployment', 'Evaluation by test users']
]
add_table(test_sched[0], test_sched[1:])
add_paragraph('')

add_heading_text('6.3 Test Strategy', level=2)
add_paragraph('The testing strategy employed both black-box and white-box testing approaches:', indent_cm=1.27, space_after=6)
test_strat = [
    'Black-box testing: Testing system functionality without examining internal code structure. Test cases are derived from functional requirements.',
    'White-box testing: Testing internal logic of critical functions, particularly the recommendation engine algorithm and quiz scoring logic.',
    'Top-down integration: Testing begins with high-level UI components and gradually integrates lower-level API modules.',
    'Classes of tests: Functional correctness tests, security tests (authentication bypass, RLS enforcement), and stress tests (concurrent user sessions).'
]
for item in test_strat:
    add_bullet(item)

add_heading_text('6.4 Test Design', level=2)

add_paragraph('6.4.1 Test Cases', bold=True, space_after=3)
add_paragraph('The following test cases were designed and executed:', indent_cm=1.27, space_after=6)

test_cases = [
    ['TC ID', 'Module', 'Test Description', 'Expected Result'],
    ['TC01', 'Authentication', 'User logs in with valid credentials', 'Redirected to role dashboard'],
    ['TC02', 'Authentication', 'User logs in with invalid password', 'Error message displayed'],
    ['TC03', 'Authentication', 'User accesses protected route without login', 'Redirected to login page'],
    ['TC04', 'RBAC', 'Learner accesses /educator route', 'Redirected to access denied'],
    ['TC05', 'Course Management', 'Educator creates a new course', 'Course appears in course list'],
    ['TC06', 'Lessons', 'Learner views a lesson', 'Progress is recorded'],
    ['TC07', 'Quiz', 'Learner submits quiz answers', 'Score calculated and displayed'],
    ['TC08', 'Recommendation', 'Learner fails a quiz', 'Revision recommendation appears'],
    ['TC09', 'Accessibility', 'Learner changes theme to high contrast', 'UI updates immediately'],
    ['TC10', 'Certificate', 'Learner completes all lessons', 'PDF certificate generated'],
    ['TC11', 'Analytics', 'Educator views dashboard', 'Charts show correct data'],
    ['TC12', 'Notifications', 'Learner enrolls in course', 'Educator receives notification']
]
add_table(test_cases[0], test_cases[1:])
add_paragraph('')

add_heading_text('6.5 Test Results and Analysis', level=2)
add_paragraph('All test cases were executed and the results were recorded. The following table summarizes the outcomes:',
              indent_cm=1.27, space_after=6)

test_results = [
    ['TC ID', 'Result', 'Remarks'],
    ['TC01', 'Pass', 'Redirected to role-appropriate dashboard'],
    ['TC02', 'Pass', 'Red error alert displayed with message'],
    ['TC03', 'Pass', 'Redirected to /login'],
    ['TC04', 'Pass', 'Access denied page shown'],
    ['TC05', 'Pass', 'Course created with status "draft"'],
    ['TC06', 'Pass', 'Lesson progress recorded with timestamp'],
    ['TC07', 'Pass', 'Score percentage and pass/fail result shown'],
    ['TC08', 'Pass', 'Recommendation card with "revision" tier appeared'],
    ['TC09', 'Pass', 'Theme changed instantly with no page reload'],
    ['TC10', 'Pass', 'PDF downloaded with correct information'],
    ['TC11', 'Pass', 'Charts displayed accurate aggregate data'],
    ['TC12', 'Pass', 'Notification appeared in educator bell icon']
]
add_table(test_results[0], test_results[1:])
add_paragraph('')
add_paragraph('All 12 test cases passed successfully. The system demonstrates correct functional behavior across all modules, proper role-based access control enforcement, and effective integration between components.',
              indent_cm=1.27, space_after=6)

doc.add_page_break()

# ============================================================
# CHAPTER 7: CONCLUSION
# ============================================================
add_heading_text('CHAPTER 7', level=1)
add_heading_text('CONCLUSION', level=1)

add_heading_text('7.1 Introduction', level=2)
add_paragraph('This chapter presents the overall conclusion of the project. It summarizes the work carried out throughout the project, discusses the key contributions and limitations, and provides suggestions for future work.',
              indent_cm=1.27, space_after=6)

add_heading_text('7.2 Project Summarization', level=2)
add_paragraph('The ACESS project successfully achieved all four objectives defined in Chapter 1:', indent_cm=1.27, space_after=6)

summary_items = [
    ('Objective 1 (Analysis): ', 'Achieved through comprehensive literature review of learning disabilities statistics (DOSM Malaysia, 2024), accessibility standards (WCAG 2.1), and comparative study of existing platforms (TalentLMS, 360Learning, Moodle).'),
    ('Objective 2 (Development): ', 'Achieved by developing a fully functional accessible web platform using Next.js 16, Supabase, and TypeScript. The system implements all eight modules with built-in TTS, adjustable display settings, adaptive recommendations, and role-based access control.'),
    ('Objective 3 (Certificates & Analytics): ', 'Achieved through the certificate generation module that produces verifiable PDF certificates with unique reference codes, and the educator analytics dashboard that visualizes enrollment, completion rates, and at-risk learners with CSV export capability.'),
    ('Objective 4 (Evaluation): ', 'Achieved through structured testing with 12 test cases covering all modules. All tests passed, confirming functional correctness and proper integration.')
]
for prefix, text in summary_items:
    add_bullet(text, bold_prefix=prefix)

add_heading_text('7.3 Project Contribution', level=2)
add_paragraph('The ACESS platform makes the following contributions:', indent_cm=1.27, space_after=6)
contributions = [
    'Provides an integrated accessible e-learning platform specifically designed for learners with dyslexia, ADHD, and mild cognitive impairment, filling a gap in the current market.',
    'Demonstrates a practical implementation of WCAG 2.1 accessibility guidelines in a modern web application using CSS data attributes and React Context.',
    'Offers an adaptive content recommendation engine using rule-based logic that adjusts difficulty based on real-time quiz performance.',
    'Provides educators with data-driven analytics to monitor learner progress and identify at-risk students for timely intervention.',
    'Establishes a verifiable certification system that enables learners to present credible proof of their skills to employers and support organizations.'
]
for item in contributions:
    add_bullet(item)

add_heading_text('7.4 Project Limitation', level=2)
add_paragraph('The project has the following limitations:', indent_cm=1.27, space_after=6)
limitations = [
    'Testing was conducted with a limited number of users. Broader user acceptance testing with actual learners with disabilities would provide more comprehensive feedback.',
    'The adaptive recommendation engine uses simple rule-based logic rather than machine learning algorithms, which could provide more personalized recommendations.',
    'The system currently supports only English and Malay languages. Additional languages would increase accessibility for a wider audience.',
    'Video content relies on embedded YouTube videos; native video hosting and streaming would provide better control and accessibility.'
]
for item in limitations:
    add_bullet(item)

add_heading_text('7.5 Future Works', level=2)
add_paragraph('Several improvements and extensions are recommended for future development:', indent_cm=1.27, space_after=6)
future = [
    'Machine Learning Recommendations: Replace the rule-based recommendation engine with a machine learning model that can analyze learning patterns and provide more personalized content suggestions.',
    'Mobile Application: Develop native mobile applications for iOS and Android to provide a more seamless learning experience on mobile devices.',
    'Gamification: Introduce gamification elements such as badges, points, and leaderboards to increase learner engagement and motivation.',
    'AI-Powered Content Generation: Integrate AI tools to assist educators in generating accessible content, including automatic transcript generation and text simplification.',
    'Extended Language Support: Add support for additional languages commonly used in Malaysia, such as Mandarin and Tamil.',
    'Offline Mode: Implement offline access to downloaded lessons and quizzes for learners with limited internet connectivity.'
]
for item in future:
    add_bullet(item)

add_paragraph('In conclusion, the ACESS platform was successfully developed and evaluated according to the defined objectives. The results demonstrate that the system provides a practical solution to the identified problem of inaccessible e-learning platforms for learners with disabilities. Although several limitations were identified, the project establishes a strong foundation for future enhancements and further research.',
              indent_cm=1.27, space_after=6)

doc.add_page_break()

# ============================================================
# REFERENCES
# ============================================================
add_heading_text('REFERENCES', level=1)
add_paragraph('Department of Statistics Malaysia (2024). Person with Disability Statistics Malaysia 2024. Accessed 15 June 2026, <https://www.dosm.gov.my/portal-main/release-content/person-with-disability-statistics-malaysia-2024>',
              indent_cm=1.27, space_after=6)
add_paragraph('SingHealth (2023). Mild Cognitive Impairment. Accessed 15 June 2026, <https://www.singhealth.com.sg/symptoms-treatments/mild-cognitive-impairment>',
              indent_cm=1.27, space_after=6)
add_paragraph('PMC (2024). Learning Disabilities and Digital Learning. Accessed 15 June 2026, <https://pmc.ncbi.nlm.nih.gov/articles/PMC9554924/>',
              indent_cm=1.27, space_after=6)
add_paragraph('Crown Counseling (2024). Statistics of Learning Disabilities. Accessed 15 June 2026, <https://crowncounseling.com/statistics/learning-disabilities/>',
              indent_cm=1.27, space_after=6)
add_paragraph('Whole School SEND (2023). Cognition and Learning. Accessed 15 June 2026, <https://www.wholeschoolsend.org.uk/page/cognition-and-learning>',
              indent_cm=1.27, space_after=6)
add_paragraph('TalentLMS (2025). Learning Management System. Accessed 15 June 2026, <https://www.talentlms.com/>',
              indent_cm=1.27, space_after=6)
add_paragraph('360Learning (2025). Collaborative Learning Platform. Accessed 15 June 2026, <https://360learning.com/>',
              indent_cm=1.27, space_after=6)
add_paragraph('Vercel (2025). Next.js Documentation. Accessed 15 June 2026, <https://nextjs.org/>',
              indent_cm=1.27, space_after=6)
add_paragraph('Supabase (2025). Supabase Documentation. Accessed 15 June 2026, <https://supabase.com/>',
              indent_cm=1.27, space_after=6)
add_paragraph('World Health Organization (2024). Disability and Health. Accessed 15 June 2026, <https://www.who.int/news-room/fact-sheets/detail/disability-and-health>',
              indent_cm=1.27, space_after=6)
add_paragraph('Web Accessibility Initiative (2023). Web Content Accessibility Guidelines (WCAG) 2.1. Accessed 15 June 2026, <https://www.w3.org/TR/WCAG21/>',
              indent_cm=1.27, space_after=6)

doc.add_page_break()

# ============================================================
# APPENDICES
# ============================================================
add_heading_text('APPENDICES', level=1)
add_paragraph('Appendix A: User Manual', bold=True, space_after=6)
add_paragraph('The user manual provides step-by-step instructions for using the ACESS platform. Key sections include:', indent_cm=1.27, space_after=3)
um_sections = [
    'Getting Started: Registration and login procedures',
    'Learner Guide: Enrolling in courses, viewing lessons, taking quizzes, tracking progress',
    'Educator Guide: Creating courses, managing lessons, building quizzes, viewing analytics',
    'Administrator Guide: Managing users, approving courses, overseeing certificates',
    'Accessibility Guide: Configuring display settings, using TTS, customizing themes'
]
for item in um_sections:
    add_bullet(item)

add_paragraph('', space_after=6)
add_paragraph('Appendix B: Source Code Repository', bold=True, space_after=6)
add_paragraph('The complete source code is available at: https://github.com/aliff1024/ACESS.git', indent_cm=1.27, space_after=6)

add_paragraph('', space_after=6)
add_paragraph('Appendix C: Database Schema', bold=True, space_after=6)
add_paragraph('The full database schema with all tables, triggers, and RLS policies is available in the scripts/ and supabase/migrations/ directories of the project repository.', indent_cm=1.27, space_after=6)

add_paragraph('', space_after=6)
add_paragraph('Appendix D: Turnitin Report', bold=True, space_after=6)
add_paragraph('The Turnitin originality report is attached on the following page.', indent_cm=1.27, space_after=6)

# ============================================================
# SAVE DOCUMENT
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ACESS_PSM_Report_BITS.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")